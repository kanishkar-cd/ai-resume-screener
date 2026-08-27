from pathlib import Path

from docx import Document

from app.schemas.parsed_document import ParserEngine
from app.services.parsers.base import ParseOutput


def parse_docx(path: Path) -> ParseOutput:
    document = Document(str(path))
    text_parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text and paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if row_cells:
                # Use dict.fromkeys to deduplicate merged cell text in table rows
                text_parts.append(" | ".join(dict.fromkeys(row_cells)))

    raw_text = "\n".join(text_parts).strip()
    return ParseOutput(
        raw_text=raw_text,
        page_count=None,
        parser_engine=ParserEngine.PYTHON_DOCX,
    )
